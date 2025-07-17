
import sys
import io
import os
import fitz  # PyMuPDF
import pandas as pd
import google.generativeai as genai
from io import StringIO
from docx import Document

# --- Fonction extraction texte depuis PDF ---
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text

# --- Fonction extraction texte depuis DOCX ---
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = ""

    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text.strip() + "\n"

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_text += "\t".join(row_data) + "\n"

    for section in doc.sections:
        header = section.header
        footer = section.footer
        if header:
            for para in header.paragraphs:
                if para.text.strip():
                    full_text += "[HEADER] " + para.text.strip() + "\n"
        if footer:
            for para in footer.paragraphs:
                if para.text.strip():
                    full_text += "[FOOTER] " + para.text.strip() + "\n"

    return full_text

# --- Fonction lecture matrice Excel ---
def read_excel_matrix(excel_path):
    return pd.read_excel(excel_path)

# --- Construction du prompt ---
def create_prompt_general(matrice_json, texte_pdf):
    prompt = f"""
Tu es un assistant intelligent spécialisé dans l’analyse de rapports statistiques.

Voici une matrice de données incomplète extraite d’un fichier Excel. Elle contient des lignes avec différents indicateurs (exemple : taux de chômage...), et des colonnes décrivant des dimensions comme l'année, la région, etc.

Chaque ligne représente un cas unique à compléter.

Voici la matrice :

{matrice_json}

Et voici le contenu textuel extrait de plusieurs fichiers PDF ou Word non structurés :

{texte_pdf}

Ta tâche est de :
1. Comprendre la structure de la matrice (colonnes et lignes).
2. Rechercher dans le texte les données correspondant à chaque ligne.
3. Remplir chaque cellule vide si la valeur peut être trouvée dans le texte.
4. Si une valeur est absente du texte ou incertaine, laisse la cellule vide.

Réponds uniquement au format CSV, avec la même structure que la matrice d’origine, remplie avec les valeurs complétées si disponibles.
**Utilise le point-virgule (;) comme séparateur de colonnes dans le CSV.**
Ne mets aucune explication, uniquement le contenu CSV.
"""
    return prompt

# --- Appel à Gemini ---
def call_gemini_api(prompt):
    genai.configure(api_key="AIzaSyBLQfllqi4UiLpjt6HUqUCQ0iBMykvWelc")  # Remplace par ta clé
    model = genai.GenerativeModel(model_name="gemini-2.5-pro")
    response = model.generate_content(prompt)
    return response.text

# --- Parseur CSV Gemini ---
def parse_gemini_csv_response(response_text):
    try:
        df = pd.read_csv(io.StringIO(response_text), sep=';', quotechar='"', encoding='utf-8')
        df.columns = df.columns.str.strip()
        df = df.apply(lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x))

        return df
    except Exception as e:
        print("Erreur lors du parsing CSV :", e)
        print("Réponse brute :", response_text)
        return pd.DataFrame()

# --- Fonction principale ---
def main(doc_type, doc_paths_csv, excel_path, output_path):
    doc_paths = doc_paths_csv.split(",")
    df = pd.read_excel(excel_path)

    full_text = ""
    for path in doc_paths:
        if doc_type == "pdf":
            full_text += extract_text_from_pdf(path) + "\n"
        else:
            full_text += extract_text_from_docx(path) + "\n"

    prompt = create_prompt_general(df.to_csv(index=False), full_text)
    response = call_gemini_api(prompt)
    df_result = parse_gemini_csv_response(response)

    if df_result.empty:
        print("Aucune donnée extraite.")
        sys.exit(1)

    df_result.to_excel(output_path, index=False)
    print("Traitement terminé.")

# --- Entrée script ---
if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python traitement.py <doc_type> <doc_paths_csv> <excel_path> <output_path>")
        sys.exit(1)

    doc_type = sys.argv[1]         # "pdf" ou "word"
    doc_paths_csv = sys.argv[2]    # chemins séparés par virgule
    excel_path = sys.argv[3]
    output_path = sys.argv[4]

    main(doc_type, doc_paths_csv, excel_path, output_path)
